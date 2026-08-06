"""
build_tables.py — Emit Tables 3–13, B1, 4e, and A1–A4 in .csv, .tex (booktabs), .docx.

Usage:
    python -m src.reporting.build_tables --config config/config.yaml

Input:  outputs/tables/raw_results.json
        outputs/tables/encompassing_results.json
        outputs/tables/degeneracy_flags.json
        outputs/tables/gate_correspondence_raw.json
        outputs/tables/ablation_ladder_raw.json
        outputs/models/<model>/<series>/{params.json, fit_info.json}
Output: outputs/tables/  *.csv  *.tex  *.docx

Table numbering (paper convention; every number below is unique --
Section 9.6 fixed a "Table 4" collision between the per-series OOS
tables and the GARCH estimation table, which is Table 4e)
------------------------------------------------------------------
Table 3       — Model roster (static)
Tables 4–7    — OOS performance, one per series (includes Panel D, the
                constant-forecast reference row, and a DEGENERATE column)
Table 4e      — GARCH(1,1) estimation (all four series)
Table 8       — Cross-market Δ% summary
Table 9       — Diebold–Mariano (+ TOST equivalence p-value column)
Table 10      — Forecast-Encompassing
Table 11a–11d — Risk backtests (Kupiec + Christoffersen + Acerbi-Szekely ES), one per series
Table 12      — Risk backtests disaggregated by market (Section 9.5;
                never averaged across series -- Fisher's method for any
                combined figure)
Table 13      — LSTM gate <-> GARCH parameter correspondence (Section 9.3)
Tables A1–A4  — Full estimation per series (all econometric models)
Table B1      — Ablation ladder / Proposition 2 verification (Section 7)
"""
from __future__ import annotations

import json
import logging
import warnings
from pathlib import Path
from typing import Any

import numpy as np
import pandas as pd
import yaml
from scipy import stats

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
log = logging.getLogger(__name__)


# ──────────────────────────────────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────────────────────────────────

def load_config(path: str) -> dict:
    with open(path) as fh:
        return yaml.safe_load(fh)


def _fmt(v, decimals=4, pct=False, bold=False) -> str:
    """Format a scalar for display."""
    if v is None or (isinstance(v, float) and not np.isfinite(v)):
        return "—"
    if pct:
        s = f"{v:+.2f}\\%"
    else:
        s = f"{v:.{decimals}f}"
    if bold:
        s = f"\\textbf{{{s}}}"
    return s


def _sig_stars(p) -> str:
    if p is None:
        return ""
    if p < 0.01:
        return "***"
    if p < 0.05:
        return "**"
    if p < 0.10:
        return "*"
    return ""


def _save_csv(df: pd.DataFrame, path: Path, panel_rows: set | None = None) -> None:
    if panel_rows:
        # A panel-divider row in CSV: keep the panel label as the index,
        # blank every data column (rather than repeating "--- Panel A ---"
        # as a literal cell value in the first data column).
        df = df.copy()
        for idx in panel_rows:
            if idx in df.index:
                df.loc[idx, :] = ""
    df.to_csv(path)
    log.info("Saved CSV: %s", path)


def _save_tex(df: pd.DataFrame, path: Path, caption: str, label: str, note: str = "",
              panel_rows: set | None = None) -> None:
    """
    Emit a booktabs LaTeX table. Rows whose index is in `panel_rows` are
    rendered as a single \\multicolumn spanning divider (e.g. "Panel A"),
    not as "Panel A & --- Panel A --- & -- & ..." repeated across every
    column (Section 9.6).

    A MultiIndex (e.g. ["Series", "Rung"]) is exploded into one leading
    "l" column per level -- str(idx) on a MultiIndex row is a Python
    tuple, which would otherwise render literally as
    "('BTC-USD', '0 -- GARCH...')" in the output.
    """
    panel_rows = panel_rows or set()
    is_multi = isinstance(df.index, pd.MultiIndex)
    idx_names = [str(n) for n in df.index.names] if is_multi else [df.index.name or ""]
    n_idx   = len(idx_names)
    n_cols  = len(df.columns)
    col_fmt = "l" * n_idx + "r" * n_cols
    lines   = [
        "\\begin{table}[ht]",
        "\\centering",
        f"\\caption{{{caption}}}",
        f"\\label{{{label}}}",
        f"\\begin{{tabular}}{{{col_fmt}}}",
        "\\toprule",
    ]
    # Header
    header  = " & ".join(idx_names) + " & " + " & ".join(str(c) for c in df.columns) + " \\\\"
    lines.append(header)
    lines.append("\\midrule")
    # Rows
    for idx, row in df.iterrows():
        idx_vals = list(idx) if is_multi else [idx]
        if idx in panel_rows:
            lines.append(f"\\multicolumn{{{n_idx + n_cols}}}{{l}}{{\\textbf{{{idx}}}}} \\\\")
            continue
        cells = " & ".join(str(v) for v in idx_vals) + " & " + " & ".join(
            str(v) if v is not None else "—" for v in row
        ) + " \\\\"
        lines.append(cells)
    lines += [
        "\\bottomrule",
        "\\end{tabular}",
    ]
    if note:
        lines.append(f"\\\\[3pt]\\footnotesize\\textit{{Note:}} {note}")
    lines.append("\\end{table}")

    path.write_text("\n".join(lines))
    log.info("Saved TEX: %s", path)


def _save_docx(df: pd.DataFrame, path: Path, title: str, note: str = "",
               panel_rows: set | None = None) -> None:
    """
    Emit a three-line MDPI-style .docx table. Panel-divider rows (Section
    9.6) get their cells merged into one bold, spanning label instead of
    repeating the panel name across every column.

    A MultiIndex (e.g. ["Series", "Rung"]) is exploded into one leading
    column per level -- str(idx) on a MultiIndex row is a Python tuple,
    which would otherwise render literally as
    "('BTC-USD', '0 -- GARCH...')" in the first cell.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        log.warning("python-docx not available; skipping .docx for %s", path)
        return

    panel_rows = panel_rows or set()
    is_multi = isinstance(df.index, pd.MultiIndex)
    idx_names = [str(n) if n else "Model" for n in df.index.names] if is_multi else [df.index.name or "Model"]
    n_idx = len(idx_names)

    doc   = Document()
    doc.add_heading(title, level=2)

    n_cols = len(df.columns) + n_idx
    table  = doc.add_table(rows=1 + len(df), cols=n_cols)
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0].cells
    for i, name in enumerate(idx_names):
        hdr[i].text = name
    for j, col in enumerate(df.columns):
        hdr[n_idx + j].text = str(col)

    # Data rows
    for i, (idx, row) in enumerate(df.iterrows()):
        cells = table.rows[i + 1].cells
        idx_vals = list(idx) if is_multi else [idx]
        if idx in panel_rows:
            merged = cells[0].merge(cells[-1]) if n_cols > 1 else cells[0]
            merged.text = str(idx)
            for para in merged.paragraphs:
                for run in para.runs:
                    run.font.bold = True
            continue
        for k, v in enumerate(idx_vals):
            cells[k].text = str(v)
        for j, v in enumerate(row):
            cells[n_idx + j].text = str(v) if v is not None else "—"

    # Style: Times New Roman 10pt
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)

    if note:
        p = doc.add_paragraph()
        run = p.add_run(f"Note: {note}")
        run.font.size = Pt(9)
        run.font.italic = True

    doc.save(str(path))
    log.info("Saved DOCX: %s", path)


def _save_all(df: pd.DataFrame, stem: Path, caption: str, label: str, note: str,
               panel_rows: set | None = None) -> None:
    """Save CSV + TEX + DOCX. `panel_rows`: index values that are panel-divider rows (Section 9.6)."""
    _save_csv(df, stem.with_suffix(".csv"), panel_rows)
    _save_tex(df, stem.with_suffix(".tex"), caption, label, note, panel_rows)
    _save_docx(df, stem.with_suffix(".docx"), caption, note, panel_rows)


# ──────────────────────────────────────────────────────────────────────────────
# Table 3 — Model Roster (static)
# ──────────────────────────────────────────────────────────────────────────────

ROSTER_DATA = [
    # Panel A — Econometric
    ("Panel A", "ARCH(1)",         "Econometric", "ARCH(1)-t",            "arch (Python)",    "Benchmark"),
    ("Panel A", "GARCH(1,1)",      "Econometric", "GARCH(1,1)-t",         "arch (Python)",    "Reference benchmark"),
    ("Panel A", "EGARCH(1,1)",     "Econometric", "EGARCH(1,1)-t",        "arch (Python)",    "Leverage effect"),
    ("Panel A", "GJR-GARCH(1,1)",  "Econometric", "GJR-GARCH(1,1)-t",     "arch (Python)",    "Threshold asymmetry"),
    ("Panel A", "FIGARCH(1,d,1)",  "Econometric", "FIGARCH(1,d,1)-t",     "arch (Python)",    "Long memory"),
    ("Panel A", "MSGARCH(1,1)",    "Econometric", "Markov-switching (2 states)", "MSGARCH (R)",  "Markov-switching (rev. req.)"),
    ("Panel A", "HAR",             "Econometric", "HAR-RV (OLS+HAC)",     "statsmodels",      "Heterogeneous volatility"),
    # Panel B — ML / DL
    ("Panel B", "SVR-GARCH",       "ML",          "SVR-RBF on ε² window", "scikit-learn",     "ML baseline"),
    ("Panel B", "NN-GARCH",        "DL",          "Dense net + GARCH prior", "Keras/TF",      "Neural GARCH augmentation (rev. req.)"),
    ("Panel B", "LSTM-SSE",        "DL",          "LSTM + MSE loss",      "Keras/TF",         "Architecture ablation"),
    ("Panel B", "CNN-LSTM",        "DL",          "Conv1D + LSTM",        "Keras/TF",         "DL baseline"),
    ("Panel B", "LSTM-Attention",  "DL",          "LSTM + Bahdanau attn", "Keras/TF",         "DL baseline"),
    ("Panel B", "TCN",             "DL",          "Dilated causal Conv",  "Keras/TF",         "DL baseline"),
    ("Panel B", "Transformer",     "DL",          "Encoder-only transformer", "Keras/TF",     "DL baseline"),
    # Panel C — Proposed
    ("Panel C", "LSTM-SSE-t-Student", "DL (proposed)",
     "LSTM + (1−λ)·MSE + λ·NLL_t", "Keras/TF", "Proposed model"),
    ("Panel C", "ARCH-LSTM", "DL (diagnostic)",
     "ARCH(1)-restricted LSTM cell (5 structural constraints)", "Keras/TF",
     "Optimizer diagnostic (ARCH(1) recovery, not a proposed forecasting model)"),
    ("Panel C", "GARCH-LSTM", "DL (diagnostic)",
     "GARCH(1,1)-restricted LSTM cell (trainable forget gate, bounded "
     "persistence/mix reparametrization)", "Keras/TF",
     "Optimizer diagnostic (GARCH(1,1) recovery, not a proposed forecasting model)"),
]


def build_table3(out_dir: Path) -> None:
    cols    = ["Panel", "Model", "Family", "Specification", "Package", "Role"]
    df      = pd.DataFrame(ROSTER_DATA, columns=cols).set_index("Model")
    note    = ("Returns scaled as 100×log-retornos (pp). All parametric models use "
               "Student-t innovations. Neural models: S=10 seeds, mean ± s.d. reported. "
               "ε²_t as volatility proxy for all evaluations.")
    stem    = out_dir / "Table3_model_roster"
    _save_all(df, stem, "Table 3. Model Roster", "tab:roster", note)


# ──────────────────────────────────────────────────────────────────────────────
# Tables 4–7 — OOS Performance (one per series)
# ──────────────────────────────────────────────────────────────────────────────

PANEL_ORDER = {
    "Panel A": ["ARCH(1)", "GARCH(1,1)", "EGARCH(1,1)", "GJR-GARCH(1,1)",
                "FIGARCH(1,d,1)", "MSGARCH(1,1)", "HAR"],
    "Panel B": ["SVR-GARCH", "NN-GARCH", "LSTM-SSE", "CNN-LSTM",
                "LSTM-Attention", "TCN", "Transformer"],
    "Panel C": ["LSTM-SSE-t-Student", "ARCH-LSTM", "GARCH-LSTM"],
    # Section 9.2: minimum-bar reference forecast every other model must beat.
    "Panel D": ["Constant (unconditional variance)"],
}

PANEL_MAP = {m: p for p, ms in PANEL_ORDER.items() for m in ms}

OOS_COLS = ["MSE", "RMSE", "MAE", "R2", "QLIKE", "LL_t_OOS", "Delta_MSE", "MCS_90", "DEGENERATE"]


def _oos_row(model: str, mdata: dict, degeneracy: dict | None = None) -> dict:
    m = mdata.get("metrics", {})
    s = mdata.get("std", {})

    def _v(key):
        v = m.get(key)
        return v if v is not None else float("nan")

    mse_v  = _v("MSE")
    rmse_v = _v("RMSE")
    mae_v  = _v("MAE")
    r2_v   = _v("R2")
    ql_v   = _v("QLIKE")
    ll_v   = _v("LL_t_OOS")
    dm_v   = _v("Delta_MSE")

    mcs = mdata.get("mcs_90")
    mcs_str = "Yes" if mcs is True else ("No" if mcs is False else "—")

    # Panel B: mean ± std format. Section 9.7: a model with no seed-to-seed
    # variation by construction (e.g. SVR-GARCH, S=1) is flagged
    # "deterministic" -- ddof=1 std of a single observation is NaN, and
    # "± nan" previously leaked straight into the table.
    if s.get("deterministic"):
        mse_str = f"{mse_v:.5f} (deterministic)"
        mae_str = f"{mae_v:.5f} (deterministic)"
    else:
        std_mse = s.get("MSE_std")
        std_mae = s.get("MAE_std")
        has_std_mse = isinstance(std_mse, (int, float)) and np.isfinite(std_mse)
        has_std_mae = isinstance(std_mae, (int, float)) and np.isfinite(std_mae)
        mse_str = (f"{mse_v:.5f} ± {std_mse:.5f}" if has_std_mse else _fmt(mse_v, 5))
        mae_str = (f"{mae_v:.5f} ± {std_mae:.5f}" if has_std_mae else _fmt(mae_v, 5))

    # Section 9.1: degeneracy flag, if available.
    deg = (degeneracy or {}).get(model)
    if deg is None:
        deg_str = "—"
    else:
        deg_str = "DEGENERATE" if deg.get("degenerate") else "ok"

    return {
        "MSE":        mse_str,
        "RMSE":       _fmt(rmse_v, 5),
        "MAE":        mae_str,
        "R2":         _fmt(r2_v, 4),
        "QLIKE":      _fmt(ql_v, 5),
        "LL_t_OOS":   _fmt(ll_v, 4),
        # Delta_MSE is relative to GARCH(1,1) -- "--" only for that one
        # reference row itself (where it would trivially be 0.00%).
        # Previously blanked any model whose NAME merely contained the
        # substring "GARCH" (EGARCH, GJR-GARCH, FIGARCH, SVR-GARCH,
        # NN-GARCH all matched), silently hiding a real, already-computed
        # Delta_MSE (Table 8 had it) from Tables 4-7.
        "Delta_MSE":  _fmt(dm_v, 2, pct=True) if model != "GARCH(1,1)" else "—",
        "MCS_90":     mcs_str,
        "DEGENERATE": deg_str,
    }


def _best_col_mask(df_raw: pd.DataFrame, col: str) -> pd.Index:
    """Return index of best (min for most, max for R2/LL) value in a column."""
    try:
        vals = pd.to_numeric(df_raw[col].apply(
            lambda x: x.split("±")[0].strip().replace("%", "").replace("\\%", "").replace("+", "")
            if isinstance(x, str) else x
        ), errors="coerce")
        if col in ("R2", "LL_t_OOS"):
            return df_raw.index[vals == vals.max()]
        else:
            return df_raw.index[vals == vals.min()]
    except Exception:
        return pd.Index([])


def build_table_oos(series: str, results: dict, tnum: int, out_dir: Path,
                     degeneracy: dict | None = None) -> None:
    rows = []
    panel_rows: set = set()
    for panel, models in PANEL_ORDER.items():
        rows.append((panel, {}))   # panel-divider row (Section 9.6: rendered as \multicolumn, not repeated text)
        panel_rows.add(panel)
        for m in models:
            if m in results:
                rows.append((m, _oos_row(m, results[m], degeneracy)))

    df_raw = pd.DataFrame(
        [r for _, r in rows],
        index=[n for n, _ in rows],
    ).fillna("—")

    # Bold best values per column (skip panel-divider rows)
    for col in ["MSE", "RMSE", "MAE", "R2", "QLIKE", "LL_t_OOS"]:
        if col not in df_raw.columns:
            continue
        best_idx = _best_col_mask(df_raw, col)
        for idx in best_idx:
            if idx in panel_rows:
                continue
            v = df_raw.at[idx, col]
            if isinstance(v, str):
                df_raw.at[idx, col] = f"\\textbf{{{v}}}"

    note = (
        f"OOS window: {results.get(list(results)[0], {}).get('n_test', '?')} observations. "
        "Returns in 100×log-ret (pp); proxy ε²_t = centred squared return. "
        "QLIKE = T⁻¹Σ[lnσ̂²_t + ε²_t/σ̂²_t]. "
        "LL_t(OOS): per-obs avg Student-t log-likelihood evaluated OOS (all models). "
        "Δ%%MSE relative to GARCH(1,1); negative = improvement. "
        "MCS 90%% based on block-bootstrap QLIKE (B=10 000, block=20). "
        "Panel B: mean ± s.d. over S=10 seeds. "
        "Panel D: fixed at the training-split unconditional variance -- the "
        "minimum bar every other model should beat. "
        "DEGENERATE (src.eval.degeneracy): temporal coefficient of variation "
        "of the OOS σ̂²_t path < 5%%, or MSE no better than Panel D's constant "
        "forecast. "
        "Standard errors Bollerslev–Wooldridge; *** p<0.01."
    )

    stem = out_dir / f"Table{tnum}_OOS_{series}"
    _save_all(
        df_raw, stem,
        f"Table {tnum}. OOS Forecasting Performance — {series}",
        f"tab:oos_{series.lower()}",
        note,
        panel_rows=panel_rows,
    )


# ──────────────────────────────────────────────────────────────────────────────
# Table 8 — Cross-market Δ% summary
# ──────────────────────────────────────────────────────────────────────────────

def build_table8(all_results: dict, series_list: list[str], out_dir: Path) -> None:
    all_models = sorted({m for s in all_results.values() for m in s})
    # Multi-level columns: series × {Delta_MSE, Delta_MAE}
    col_tuples = [(s, metric) for s in series_list for metric in ["Δ%MSE", "Δ%MAE"]]
    cols       = pd.MultiIndex.from_tuples(col_tuples)

    rows = []
    for model in all_models:
        row = []
        for series in series_list:
            m_res = all_results.get(series, {}).get(model, {})
            metrics = m_res.get("metrics", {})
            dm_holm  = m_res.get("dm_qlike", {})
            reject   = dm_holm.get("reject", False)
            dm_stat  = dm_holm.get("DM_stat")

            dm_pct  = metrics.get("Delta_MSE")
            mae_pct = metrics.get("Delta_MAE")

            def _cell(v):
                # GARCH(1,1) is the reference every Delta_% is computed
                # against, so its own row is trivially 0.00% -- shown as
                # "--", not "+0.00" (previously this special case was
                # documented in the note but never implemented).
                if model == "GARCH(1,1)":
                    return "—"
                if v is None or (isinstance(v, float) and not np.isfinite(v)):
                    return "—"
                s = f"{v:+.2f}"
                if not reject or dm_stat is None:
                    return s
                # Section 9.6: two distinct markers by WHO wins the DM test
                # (src.eval.dm_test: positive DM_stat -> proposed wins).
                # Previously this bolded the rival's cell whenever DM was
                # significant AT ALL, including when the rival won --
                # reading as the opposite of what bold is supposed to mean.
                if dm_stat > 0:
                    return f"\\textbf{{{s}}}"     # proposed significantly better than this rival
                else:
                    return f"\\underline{{{s}}}"  # rival significantly better than proposed

            row.extend([_cell(dm_pct), _cell(mae_pct)])
        rows.append(row)

    df = pd.DataFrame(rows, index=all_models, columns=cols)
    note = (
        "Percentage change relative to GARCH(1,1) (reference row shown as "
        "--). \\textbf{Bold} = proposed model significantly beats this "
        "rival (positive DM stat, QLIKE, Holm-corrected p<0.05); "
        "\\underline{underlined} = this rival significantly beats the "
        "proposed model (negative DM stat, same test) -- these are "
        "opposite conclusions and were previously conflated under a single "
        "bold marker. Unmarked = not statistically significant. "
        "Negative Delta_%% = improvement over GARCH(1,1)."
    )
    stem = out_dir / "Table8_CrossMarket_Summary"
    _save_all(df, stem, "Table 8. Cross-Market OOS Performance Summary",
              "tab:cross_market", note)


# ──────────────────────────────────────────────────────────────────────────────
# Table 9 — Diebold–Mariano
# ──────────────────────────────────────────────────────────────────────────────

def build_table9(all_results: dict, series_list: list[str], out_dir: Path) -> None:
    all_models = sorted({
        m for s in all_results.values()
        for m in s if m != "LSTM-SSE-t-Student"
    })

    col_tuples = [(s, stat) for s in series_list for stat in ["DM", "p(Holm)", "TOST p"]]
    cols = pd.MultiIndex.from_tuples(col_tuples)

    rows = []
    for rival in all_models:
        row = []
        for series in series_list:
            model_res = all_results.get(series, {}).get(rival, {})
            dm_res = model_res.get("dm_qlike", {})
            dm_v   = dm_res.get("DM_stat")
            p_holm = dm_res.get("p_holm")
            stars  = _sig_stars(p_holm)
            dm_str = f"{dm_v:.3f}{stars}" if dm_v is not None else "—"
            ph_str = f"{p_holm:.3f}" if p_holm is not None else "—"

            tost_res = model_res.get("tost", {})
            p_tost = tost_res.get("p_tost")
            if p_tost is None:
                tost_str = "—"
            else:
                tost_str = f"{p_tost:.3f}" + ("\\textsuperscript{eq}" if tost_res.get("equivalent") else "")
            row.extend([dm_str, ph_str, tost_str])
        rows.append(row)

    df = pd.DataFrame(rows, index=all_models, columns=cols)
    note = (
        "Diebold–Mariano test: LSTM-SSE-t-Student vs. each rival model. "
        "Loss function: QLIKE. HAC SE (Newey–West, bandwidth ⌊4(T/100)^{2/9}⌋). "
        "p(Holm): p-value after Holm–Bonferroni FWER correction per market. "
        "Positive DM stat: rival has higher loss → proposed wins. "
        "Significance: *** p<0.01, ** p<0.05, * p<0.10 (after Holm). "
        "TOST p (Section 9.4): two one-sided-tests p-value for equivalence "
        "of QLIKE loss within margin delta = tost.delta\\_pct × GARCH(1,1)'s "
        "own mean QLIKE (config, default 2%%), HAC SE. TOST p < 0.05 "
        "(marked \\textsuperscript{eq}) is a POSITIVE claim of equivalence -- "
        "the opposite reading from a DM p-value."
    )
    stem = out_dir / "Table9_DieboldMariano"
    _save_all(df, stem, "Table 9. Diebold–Mariano Tests", "tab:dm", note)


# ──────────────────────────────────────────────────────────────────────────────
# Table 10 — Forecast-Encompassing Tests
# ──────────────────────────────────────────────────────────────────────────────

def build_table10(encompassing_all: dict, series_list: list[str], out_dir: Path) -> None:
    """
    ε²_t = b0 + b1·σ²_benchmark + b2·σ²_candidate + u_t   (HAC SE).

    One row per (candidate, benchmark) pair; columns show β_candidate
    (with significance stars on p_candidate) and the verdict per series.
    """
    all_pairs = sorted({
        pair for s in encompassing_all.values() for pair in s
    })

    col_tuples = [(s, stat) for s in series_list for stat in ["β_cand", "verdict"]]
    cols = pd.MultiIndex.from_tuples(col_tuples)

    verdict_labels = {
        "candidate_adds_info":  "Adds info",
        "benchmark_sufficient": "Benchmark suff.",
        "both_contribute":      "Both contribute",
        "neither_significant":  "Neither sig.",
    }

    rows = []
    for pair in all_pairs:
        row = []
        for series in series_list:
            res = encompassing_all.get(series, {}).get(pair, {})
            if "error" in res:
                row.extend(["—", "n/a"])
                continue
            b_cand  = res.get("beta_candidate")
            p_cand  = res.get("p_candidate")
            stars   = _sig_stars(p_cand)
            b_str   = f"{b_cand:.3f}{stars}" if b_cand is not None else "—"
            verdict = verdict_labels.get(res.get("verdict"), "—")
            row.extend([b_str, verdict])
        rows.append(row)

    pair_labels = [p.replace("_vs_", " vs. ") for p in all_pairs]
    df = pd.DataFrame(rows, index=pair_labels, columns=cols)

    note = (
        "Forecast-encompassing regression: \\varepsilon^2_t = b_0 + b_1 "
        "\\sigma^2_{benchmark,t} + b_2 \\sigma^2_{candidate,t} + u_t, HAC SE "
        "(Newey--West, bandwidth \\lfloor 4(T/100)^{2/9} \\rfloor). "
        "\\beta_{cand}: coefficient on the candidate forecast (significance "
        "from its own HAC p-value). 'Adds info': candidate significant, "
        "benchmark not — benchmark is redundant given the candidate. "
        "'Both contribute': both significant — complementary information. "
        "'Benchmark suff.': candidate carries no information beyond the "
        "benchmark. 'Neither sig.': inconclusive (often collinearity). "
        "Significance: *** p<0.01, ** p<0.05, * p<0.10."
    )
    stem = out_dir / "Table10_Encompassing"
    _save_all(df, stem, "Table 10. Forecast-Encompassing Tests", "tab:encompassing", note)


# ──────────────────────────────────────────────────────────────────────────────
# Table 11a–11d — Risk backtests (Kupiec + Christoffersen + VaR + ES)
# ──────────────────────────────────────────────────────────────────────────────

def _fmt_bool_pass(p_uc: Any) -> str:
    if p_uc is None or (isinstance(p_uc, float) and not np.isfinite(p_uc)):
        return "—"
    return "Yes" if p_uc >= 0.05 else "No"


def _fmt_bool_pass_cc(p_cc: Any) -> str:
    if p_cc is None or (isinstance(p_cc, float) and not np.isfinite(p_cc)):
        return "—"
    return "Yes" if p_cc >= 0.05 else "No"


def _n_expected(st: dict, level: str) -> str:
    T = st.get("T")
    if T is None:
        return "—"
    alpha = 1.0 - float(level)
    return f"{T * alpha:.2f}"


def _risk_row(m_res: dict, level: str) -> dict:
    ves = m_res.get("var_es", {}).get(level, {})
    st  = ves.get("student_t", {})
    esbt = ves.get("es_backtest", {})
    return {
        f"n_exc@{level}": str(st.get("n_exc", "—")),
        f"n_exp@{level}": _n_expected(st, level),
        f"ExcRate@{level}": _fmt(st.get("exc_rate"), decimals=4),
        f"LR_uc@{level}": _fmt(st.get("LR_uc"), decimals=4),
        f"p_uc@{level}": _fmt(st.get("p_uc"), decimals=4),
        f"KupiecPass@{level}": _fmt_bool_pass(st.get("p_uc")),
        f"LR_ind@{level}": _fmt(st.get("LR_ind"), decimals=4),
        f"p_ind@{level}": _fmt(st.get("p_ind"), decimals=4),
        f"LR_cc@{level}": _fmt(st.get("LR_cc"), decimals=4),
        f"p_cc@{level}": _fmt(st.get("p_cc"), decimals=4),
        f"ChristoffersenPass@{level}": _fmt_bool_pass_cc(st.get("p_cc")),
        f"VaR_t_mean@{level}": _fmt(ves.get("var_t_mean"), decimals=6),
        f"ES_t_mean@{level}": _fmt(ves.get("es_t_mean"), decimals=6),
        f"ES_Z2@{level}": _fmt(esbt.get("Z2"), decimals=4),
        f"ES_p@{level}": _fmt(esbt.get("p_value"), decimals=4),
    }


def build_table11_risk(series: str, results: dict, tlabel: str, out_dir: Path) -> None:
    if not results:
        return

    all_models = sorted(results.keys())
    levels = sorted({
        lv
        for m in results.values()
        for lv in m.get("var_es", {}).keys()
    }, key=float, reverse=True)

    rows = []
    for model in all_models:
        m_res = results.get(model, {})
        row = {}
        for lv in levels:
            row.update(_risk_row(m_res, lv))
        rows.append((model, row))

    if not rows:
        return

    df = pd.DataFrame([r for _, r in rows], index=[m for m, _ in rows])
    note = (
        "Student-t VaR backtest over OOS residuals. n_exc/n_exp: observed vs. "
        "expected exceedances (T x tail probability). Kupiec uses the UC LR test "
        "for correct unconditional exceedance rate (H0: hit rate = tail probability). "
        "Christoffersen uses the CC LR test (UC + independence). "
        "KupiecPass = Yes when p_uc >= 0.05; ChristoffersenPass = Yes when p_cc >= 0.05. "
        "ES_Z2/ES_p (Section 9.5): Acerbi-Szekely (2014) Z2 Expected-Shortfall "
        "backtest, HAC-free Monte Carlo p-value under the model's own posited "
        "Student-t distribution (H0: ES correctly specified; not just an "
        "average magnitude). VaR_t_mean/ES_t_mean are averages of model-implied "
        "risk forecasts over the OOS window."
    )
    stem = out_dir / f"Table11_{tlabel}_Risk_{series}"
    _save_all(
        df,
        stem,
        f"Table 11{tlabel}. Risk Backtest Summary — {series}",
        f"tab:risk_{series.lower()}",
        note,
    )


def _fisher_combine(p_values: list) -> tuple[float, float]:
    """
    Fisher's method for combining independent p-values:
        X2 = -2 * sum(ln(p_i)) ~ chi2(2k) under H0 (all p_i uniform).
    Returns (X2_stat, combined_p). NEVER use the arithmetic mean of
    p-values to summarize across markets -- it has no valid null
    distribution and can badly mislead (Section 9.5).
    """
    valid = [max(float(p), 1e-300) for p in p_values
             if isinstance(p, (int, float)) and np.isfinite(p) and p > 0]
    if not valid:
        return float("nan"), float("nan")
    x2 = -2.0 * sum(np.log(p) for p in valid)
    df = 2 * len(valid)
    combined_p = float(1.0 - stats.chi2.cdf(x2, df=df))
    return float(x2), combined_p


def build_table12_risk_summary(all_results: dict, series_list: list[str], out_dir: Path) -> None:
    """
    Table 12 (fixed, Section 9.5): risk backtests disaggregated by market
    -- one row per (series, level, model), never averaged across markets.
    Ends with a "Combined (Fisher)" pseudo-series per (level, model) that
    validly combines the four markets' p-values via Fisher's method
    (replacing the previous, invalid arithmetic mean of p-values).
    """
    all_models = sorted({m for s in all_results.values() for m in s})
    levels = ["0.99", "0.975"]

    rows = []
    for lv in levels:
        for model in all_models:
            p_uc_list, p_ind_list, p_cc_list, es_p_list = [], [], [], []
            for series in series_list:
                m_res = all_results.get(series, {}).get(model, {})
                ves = m_res.get("var_es", {}).get(lv, {})
                st = ves.get("student_t", {})
                esbt = ves.get("es_backtest", {})

                rows.append({
                    "Series": series, "Level": lv, "Model": model,
                    "n_exc": st.get("n_exc"),
                    "n_exp": (float(st["T"]) * (1.0 - float(lv))) if st.get("T") is not None else None,
                    "CoverageRate": st.get("exc_rate"),
                    "LR_uc": st.get("LR_uc"), "p_uc": st.get("p_uc"),
                    "LR_ind": st.get("LR_ind"), "p_ind": st.get("p_ind"),
                    "LR_cc": st.get("LR_cc"), "p_cc": st.get("p_cc"),
                    "ES_Z2": esbt.get("Z2"), "ES_p": esbt.get("p_value"),
                })
                for lst, key, src in [(p_uc_list, "p_uc", st), (p_ind_list, "p_ind", st),
                                        (p_cc_list, "p_cc", st), (es_p_list, "p_value", esbt)]:
                    v = src.get(key)
                    if isinstance(v, (int, float)) and np.isfinite(v):
                        lst.append(v)

            _, p_uc_comb = _fisher_combine(p_uc_list)
            _, p_ind_comb = _fisher_combine(p_ind_list)
            _, p_cc_comb = _fisher_combine(p_cc_list)
            _, es_p_comb = _fisher_combine(es_p_list)
            rows.append({
                "Series": "Combined (Fisher)", "Level": lv, "Model": model,
                "n_exc": None, "n_exp": None, "CoverageRate": None,
                "LR_uc": None, "p_uc": p_uc_comb,
                "LR_ind": None, "p_ind": p_ind_comb,
                "LR_cc": None, "p_cc": p_cc_comb,
                "ES_Z2": None, "ES_p": es_p_comb,
            })

    if not rows:
        return

    df_raw = pd.DataFrame(rows).set_index(["Series", "Level", "Model"])
    df = pd.DataFrame(index=df_raw.index)
    int_cols = {"n_exc"}
    for col in df_raw.columns:
        if col in int_cols:
            df[col] = df_raw[col].map(lambda x: "—" if x is None or not np.isfinite(x) else f"{int(x)}")
        else:
            df[col] = df_raw[col].map(lambda x: "—" if x is None or not np.isfinite(x) else f"{x:.4f}")

    note = (
        "Risk backtests disaggregated by market (Section 9.5) -- never "
        "averaged across series. n_exc/n_exp: observed vs. expected "
        "exceedances. Kupiec (UC), Christoffersen independence (IND) and "
        "conditional coverage (CC) LR tests, Student-t VaR. ES_Z2/ES_p: "
        "Acerbi-Szekely (2014) Expected-Shortfall backtest (Monte Carlo "
        "p-value). 'Combined (Fisher)' rows validly combine the four "
        "markets' p-values via Fisher's method (X2 = -2*sum(ln(p_i)) ~ "
        "chi2(2k)) -- NOT their arithmetic mean, which has no valid null "
        "distribution. p >= 0.05 fails to reject the corresponding null "
        "(correct coverage / independence / correct ES) at 5%%."
    )
    stem = out_dir / "Table12_risk_by_market"
    _save_all(df, stem, "Table 12. Risk Backtests by Market", "tab:risk_by_market", note)


# ──────────────────────────────────────────────────────────────────────────────
# Table 4e — GARCH(1,1) estimation (all four series)
# ──────────────────────────────────────────────────────────────────────────────

GARCH11_FOLDERS = ["GARCH11", "GARCH_1_1_", "GARCH(1,1)"]


def _load_garch_params(models_dir: Path, series: str) -> tuple[dict, dict]:
    for folder in GARCH11_FOLDERS:
        p = models_dir / folder / series / "params.json"
        fi = models_dir / folder / series / "fit_info.json"
        if p.exists():
            params = json.loads(p.read_text())
            fit    = json.loads(fi.read_text()) if fi.exists() else {}
            return params, fit
    return {}, {}


def build_table4_estimation(
    cfg: dict,
    models_dir: Path,
    processed_dir: Path,
    out_dir: Path,
) -> None:
    series_list = [s["name"] for s in cfg["series"]]
    rows = []

    PARAM_ROWS = [
        ("Panel A: Parameters", None),
        ("ω (omega)", "omega"),
        ("α₁ (alpha[1])", "alpha[1]"),
        ("β₁ (beta[1])", "beta[1]"),
        ("ν (nu)", "nu"),
        ("Panel B: Persistence", None),
        ("α₁+β₁", "_persistence"),
        ("Half-life (days)", "_half_life"),
        ("Stationary (α+β<1)", "_stationary"),
        ("Unconditional σ² = ω/(1−α−β)", "_uncond_var"),
        ("Panel C: In-sample fit", None),
        ("LL (in-sample)", "_LL"),
        ("AIC", "_AIC"),
        ("BIC", "_BIC"),
        ("N (obs)", "_n_obs"),
        ("Panel D: Diagnostics", None),
        ("Convergence flag", "_conv"),
        ("Hessian PD", "_hess_pd"),
    ]

    for label, key in PARAM_ROWS:
        if key is None:
            row = {s: "" for s in series_list}
            row["_label"] = label
        else:
            row = {"_label": label}
            for series in series_list:
                params, fi = _load_garch_params(models_dir, series)
                if key.startswith("_"):
                    # From fit_info
                    k2 = {
                        "_persistence": "persistence",
                        "_half_life":   "half_life_days",
                        "_stationary":  "stationary",
                        "_uncond_var":  "uncond_var",
                        "_LL":          "LL_insample",
                        "_AIC":         "AIC",
                        "_BIC":         "BIC",
                        "_n_obs":       "n_obs",
                        "_conv":        "convergence",
                        "_hess_pd":     "hess_pd",
                    }.get(key, key.lstrip("_"))
                    v = fi.get(k2, None)
                else:
                    v = params.get(key, None)

                if v is None:
                    row[series] = "—"
                elif isinstance(v, bool):
                    row[series] = str(v)
                elif isinstance(v, (int, float)) and np.isfinite(v):
                    row[series] = f"{v:.4f}"
                else:
                    row[series] = str(v)
        rows.append(row)

    df = pd.DataFrame(rows).set_index("_label")
    df.index.name = "Parameter"
    panel_rows = {label for label, key in PARAM_ROWS if key is None}

    note = (
        "GARCH(1,1) with Student-t innovations estimated by MLE. "
        "Returns in 100×log-ret (pp). "
        "Standard errors: Bollerslev–Wooldridge robust (outer-product of gradient). "
        "Unconditional variance = ω/(1−α₁−β₁); compared to empirical ε² variance "
        f"(warning if deviation > {cfg.get('unconditional_var_tolerance', 0.20):.0%}). "
        "Half-life = ln(0.5)/ln(α₁+β₁). *** p<0.01, ** p<0.05, * p<0.10."
    )
    # Numbering note (Section 9.6): this is "Table 4e" (per the module
    # docstring), a distinct sub-table alongside the four per-series OOS
    # tables that are ALSO numbered 4-7 (build_table_oos). The caption
    # previously said "Table 4" verbatim, colliding with Table 4
    # (OOS Forecasting Performance -- BTC-USD) -- fixed here.
    stem = out_dir / "Table4e_GARCH11_Estimation"
    _save_all(df, stem, "Table 4e. GARCH(1,1) Parameter Estimates",
              "tab:garch_est_4e", note, panel_rows=panel_rows)


# ──────────────────────────────────────────────────────────────────────────────
# Tables A1–A4 — Full estimation by series
# ──────────────────────────────────────────────────────────────────────────────

ECON_MODEL_FOLDERS = {
    "ARCH(1)":        ["ARCH1", "ARCH_1_"],
    "GARCH(1,1)":     ["GARCH11", "GARCH_1_1_"],
    "EGARCH(1,1)":    ["EGARCH11", "EGARCH_1_1_"],
    "GJR-GARCH(1,1)": ["GJR-GARCH11", "GJR-GARCH_1_1_", "GJR_GARCH_1_1_"],
    "FIGARCH(1,d,1)": ["FIGARCH1d1", "FIGARCH_1_d_1_", "FIGARCH11d"],
    "HAR":            ["HAR"],
    "MSGARCH(1,1)":   ["MSGARCH"],
}

PARAM_KEYS_BY_MODEL = {
    "ARCH(1)":        ["omega", "alpha[1]", "nu"],
    "GARCH(1,1)":     ["omega", "alpha[1]", "beta[1]", "nu"],
    "EGARCH(1,1)":    ["omega", "alpha[1]", "gamma[1]", "beta[1]", "nu"],
    "GJR-GARCH(1,1)": ["omega", "alpha[1]", "gamma[1]", "beta[1]", "nu"],
    "FIGARCH(1,d,1)": ["omega", "phi[1]", "d", "beta[1]", "nu"],
    "HAR":            ["const", "beta_d", "beta_w", "beta_m"],
    # Two-regime switching GARCH(1,1)-t (regime.const="nu" -- shared df
    # across regimes is NOT enforced by the installed MSGARCH version,
    # each regime gets its own nu_1/nu_2). Names match fit$par exactly
    # (Section 9.7 -- confirmed by actually running R/msgarch.R; coef(fit)
    # returns NULL in this version, params previously came from nowhere).
    "MSGARCH(1,1)":   ["alpha0_1", "alpha1_1", "beta_1", "nu_1",
                        "alpha0_2", "alpha1_2", "beta_2", "nu_2",
                        "P_1_1", "P_2_1"],
}

ALL_PARAM_ROWS = [
    "omega", "alpha[1]", "gamma[1]", "beta[1]", "phi[1]", "d",
    "const", "beta_d", "beta_w", "beta_m",
    "alpha0_1", "alpha1_1", "beta_1", "nu_1",
    "alpha0_2", "alpha1_2", "beta_2", "nu_2", "P_1_1", "P_2_1",
    "nu",
    "LL_insample", "AIC", "BIC", "convergence", "hess_pd", "n_obs",
]


def _load_model_data(models_dir: Path, model_display: str, series: str) -> tuple[dict, dict]:
    for folder in ECON_MODEL_FOLDERS.get(model_display, []):
        p  = models_dir / folder / series / "params.json"
        fi = models_dir / folder / series / "fit_info.json"
        if p.exists():
            params  = json.loads(p.read_text())
            fit_inf = json.loads(fi.read_text()) if fi.exists() else {}
            return params, fit_inf
    return {}, {}


def build_table_Ax(
    series: str,
    idx: int,
    models_dir: Path,
    out_dir: Path,
) -> None:
    models = list(ECON_MODEL_FOLDERS.keys())
    rows   = []

    for param_key in ALL_PARAM_ROWS:
        row = {"_param": param_key}
        for model in models:
            params, fi = _load_model_data(models_dir, model, series)
            combined   = {**params, **fi}
            v = combined.get(param_key, None)
            if v is None:
                row[model] = "—"
            elif isinstance(v, bool):
                row[model] = str(v)
            elif isinstance(v, (int, float)) and np.isfinite(v):
                row[model] = f"{v:.4f}"
            else:
                row[model] = str(v)
        rows.append(row)

    df = pd.DataFrame(rows).set_index("_param")
    df.index.name = "Parameter"

    note = (
        f"Series: {series}. Returns in 100×log-ret (pp). All parametric models: "
        "Student-t innovations. '—' = parameter not defined in this specification. "
        "LL_insample: total log-likelihood on training sample. "
        "Standard errors: Bollerslev–Wooldridge robust. *** p<0.01."
    )
    stem = out_dir / f"TableA{idx}_Estimation_{series}"
    _save_all(
        df, stem,
        f"Table A{idx}. Parameter Estimates — {series}",
        f"tab:est_{series.lower()}",
        note,
    )


# ──────────────────────────────────────────────────────────────────────────────
# Table 13 — Gate <-> GARCH parameter correspondence (Section 9.3)
#
# Numbering note: the brief names this "Table 11", but Table 11 (a-d, risk
# backtests) and Table 12 (cross-market risk summary) already existed in
# this codebase before the respecification (per project decision: keep
# existing numbering, append brief-only new tables after it) -- so this is
# Table 13, the next free slot.
# ──────────────────────────────────────────────────────────────────────────────

def build_table13_gate_correspondence(tables_dir: Path) -> None:
    """
    Table 13: per series, E[i_t]/E[f_t] (+ sd, correlation with eps2_t),
    the alpha/beta Proposition-2 gate-mean mapping vs. GARCH-t's own MLE,
    implied persistence/half-life, and the sigma2_LSTM = a + b*sigma2_GARCH
    regression (Pearson r, b, its 95% CI, and whether that CI contains 1).

    Source: outputs/tables/gate_correspondence_raw.json
    (src.eval.gate_correspondence.run).
    """
    raw_path = tables_dir / "gate_correspondence_raw.json"
    if not raw_path.exists():
        log.warning("gate_correspondence_raw.json not found — skipping Table 13 "
                     "(run `python -m src.eval.gate_correspondence` first).")
        return
    raw = json.loads(raw_path.read_text())

    rows = []
    for series, r in raw.items():
        reg = r.get("regression") or {}
        rows.append({
            "Series": series,
            "E[i_t]": _fmt(r["E_i"]["mean"], 4),
            "sd[i_t] (across seeds)": _fmt(r["E_i"]["std"], 4),
            "E[f_t]": _fmt(r["E_f"]["mean"], 4),
            "sd[f_t] (across seeds)": _fmt(r["E_f"]["std"], 4),
            "corr(i_t, eps2_t)": _fmt(r.get("corr_i_eps2"), 4),
            "corr(f_t, eps2_t)": _fmt(r.get("corr_f_eps2"), 4),
            "alpha_implied": _fmt(r["alpha_implied"], 4),
            "alpha_hat (GARCH-t)": _fmt(r["alpha_garch"], 4),
            "beta_implied": _fmt(r["beta_implied"], 4),
            "beta_hat (GARCH-t)": _fmt(r["beta_garch"], 4),
            "Persistence (implied)": _fmt(r["persistence_implied"], 4),
            "Persistence (GARCH-t)": _fmt(r["persistence_garch"], 4),
            "Half-life implied (days)": _fmt(r["half_life_implied"], 1),
            "Half-life GARCH-t (days)": _fmt(r["half_life_garch"], 1),
            "Pearson r(LSTM,GARCH)": _fmt(reg.get("pearson_r"), 4),
            "b (slope)": _fmt(reg.get("b"), 4),
            "b 95% CI": (
                f"[{reg['b_ci_low']:.4f}, {reg['b_ci_high']:.4f}]" if reg.get("b_ci_low") is not None else "—"
            ),
            "CI contains 1": "Yes" if reg.get("b_ci_contains_1") else ("No" if reg else "—"),
        })

    df = pd.DataFrame(rows).set_index("Series")
    note = (
        "E[i_t]/E[f_t]: mean input/forget gate activation over the OOS test "
        "windows, averaged across seeds (reconstructed forward pass -- Keras "
        "does not expose intermediate gate activations). alpha\\_implied = "
        "E[i\\_t], beta\\_implied = E[f\\_t] under Proposition 2's mapping "
        "(exact only when gates are structurally constant, as in Section 7's "
        "Rung 1; sd[i\\_t]/sd[f\\_t] measure how close the free-gate trained "
        "model stays to that regime -- see src.eval.gate\\_correspondence "
        "module docstring for the full derivation). "
        "sigma2\\_LSTM = a + b*sigma2\\_GARCH (HAC SE): equivalence predicts "
        "b~=1; 'CI contains 1' is a direct, assumption-light equivalence "
        "check independent of the gate-mean mapping above."
    )
    stem = tables_dir / "Table13_gate_correspondence"
    _save_all(df, stem, "Table 13. LSTM Gate <-> GARCH Parameter Correspondence",
              "tab:gate_correspondence", note)


# ──────────────────────────────────────────────────────────────────────────────
# Table B1 — Ablation ladder (Section 7 / Proposition 2 verification)
# ──────────────────────────────────────────────────────────────────────────────

def build_table_b1(tables_dir: Path) -> None:
    """
    Table B1: per series x rung (0-3), the recovered/estimated GARCH-like
    parameters (rungs 0-1) or a cross-reference to Table 11's gate
    correspondence (rungs 2-3, where alpha/beta are not directly
    estimated parameters but gate statistics -- see
    src.eval.gate_correspondence, Section 9.3), plus implied
    persistence, half-life, QLIKE(OOS), LL_t(OOS).

    Source: outputs/tables/ablation_ladder_raw.json
    (src.models.ablation_ladder.run). Rungs 2-3 rows are left with
    alpha/beta/omega/nu = '—' if that model hasn't been trained yet, or
    if not yet cross-referenced to Table 11.
    """
    raw_path = tables_dir / "ablation_ladder_raw.json"
    if not raw_path.exists():
        log.warning("ablation_ladder_raw.json not found — skipping Table B1 "
                     "(run `python -m src.models.ablation_ladder` first).")
        return
    rows_raw = json.loads(raw_path.read_text())

    rung_labels = {
        0: "0 — GARCH(1,1)-t (MLE)",
        1: "1 — Constrained LSTM (pure Student-t MLE)",
        2: "2 — LSTM-SSE (free gates, SSE loss)",
        3: "3 — LSTM-SSE-t-Student (proposed)",
    }

    rows = []
    for r in rows_raw:
        rows.append({
            "Series": r["series"],
            "Rung": rung_labels.get(r["rung"], str(r["rung"])),
            "alpha (or E[i_t], Table 11)": _fmt(r.get("alpha"), 4),
            "beta (or E[f_t], Table 11)": _fmt(r.get("beta"), 4),
            "omega": _fmt(r.get("omega"), 4),
            "nu": _fmt(r.get("nu"), 3),
            "Persistence": _fmt(r.get("persistence"), 4),
            "Half-life (days)": _fmt(r.get("half_life_days"), 1),
            "QLIKE (OOS)": _fmt(r.get("qlike_oos"), 4),
            "LL_t (OOS)": _fmt(r.get("ll_t_oos"), 4),
            "MCS member (90%)": "—" if r.get("mcs_member") is None else str(r["mcs_member"]),
        })

    df = pd.DataFrame(rows).set_index(["Series", "Rung"])

    note = (
        "Rung 0: GARCH(1,1)-t maximum likelihood (reference). Rung 1: single-unit "
        "LSTM cell with input/forget gates held constant-but-trainable (kernels "
        "fixed at zero), output gate identically 1, initial cell state fixed, "
        "trained by pure Student-t log-likelihood (L-BFGS-B) from a neutral "
        "start (alpha0=0.05, beta0=0.85) -- NOT the GARCH answer -- so recovering "
        "alpha\\_hat, beta\\_hat is a genuine, non-circular test of Proposition 2 "
        "(see logs/proposition2\\_check.log for the PASS/FAIL verdict per "
        "series, tolerance 10\\% relative error). Rungs 2-3 report the "
        "already-trained LSTM-SSE / LSTM-SSE-t-Student models as-is; their "
        "gate statistics (analogous to alpha, beta under Proposition 2's "
        "mapping) are in Table 11, not duplicated here. Half-life = "
        "ln(0.5)/ln(beta). QLIKE and LL\\_t evaluated OOS on the test split."
    )
    stem = tables_dir / "TableB1_ablation_ladder"
    _save_all(df, stem, "Table B1. Ablation Ladder — Proposition 2 Verification",
              "tab:ablation_b1", note)


# ──────────────────────────────────────────────────────────────────────────────
# Table C2 — ARCH(1)-restricted LSTM (λ, ν) sensitivity sweep
# ──────────────────────────────────────────────────────────────────────────────

def build_table_c2_lambda_nu_sensitivity(tables_dir: Path, models_dir: Path) -> None:
    """
    Table C2: per series x (lambda, nu_fixed) grid point, the recovered
    (alpha_hat, omega_hat) vs arch's own ARCH(1)-t MLE reference, relative
    recovery error, OOS QLIKE/LL_t, and the 10%-tolerance verdict.

    Source: outputs/models/ARCH-LSTM/lambda_nu_sensitivity/
    <series>_lambda_nu_sensitivity.json, one file per series that has been
    run (src.eval.arch_restricted_recovery.run_lambda_nu_sensitivity_series
    / --lambda-nu-sensitivity). nu is held FIXED (not learned) at each grid
    point -- this sweep measures bias from misspecifying the assumed
    tail-heaviness, not whether the optimizer converges to the right nu
    (see that module's docstring). Silently includes whatever subset of
    series/grid points have been run so far; does not require the full
    config-specified grid to be complete.
    """
    pattern = "*_lambda_nu_sensitivity.json"
    src_dir = models_dir / "ARCH-LSTM" / "lambda_nu_sensitivity"
    raw_paths = sorted(src_dir.glob(pattern)) if src_dir.exists() else []
    if not raw_paths:
        log.warning("No %s files found under %s — skipping Table C2 "
                     "(run src.eval.arch_restricted_recovery --lambda-nu-sensitivity first).",
                     pattern, src_dir)
        return

    rows = []
    for p in raw_paths:
        for r in json.loads(p.read_text()):
            rows.append({
                "Series": r["series"],
                "lambda": f"{r['lam']:.2f}",
                "nu (fixed)": f"{r['nu_fixed']:g}",
                "alpha_hat": _fmt(r["alpha_recovered"], 4),
                "alpha_ref": _fmt(r["alpha_ref"], 4),
                "Rel. err. alpha": f"{100 * r['rel_err_alpha']:.2f}\\%",
                "omega_hat": _fmt(r["omega_recovered"], 4),
                "omega_ref": _fmt(r["omega_ref"], 4),
                "Rel. err. omega": f"{100 * r['rel_err_omega']:.2f}\\%",
                "QLIKE (OOS)": _fmt(r["qlike_oos"], 4),
                "LL_t (OOS)": _fmt(r["ll_t_oos"], 4),
                "Verdict (10% tol.)": r["verdict"],
            })

    df = pd.DataFrame(rows).set_index(["Series", "lambda", "nu (fixed)"])

    note = (
        "lambda: hybrid-loss mixing weight ((1-lambda)*SSE + lambda*Student-t "
        "NLL); only lambda=1.0 is apples-to-apples with arch's own MLE (see "
        "src.models.arch\\_restricted's \\_LAM\\_PURE\\_MLE docstring) -- "
        "lambda<1.0 rows measure how much recovery degrades away from that "
        "anchor. nu (fixed): Student-t degrees of freedom imposed, NOT "
        "learned (nu\\_mode=\"fixed\") -- tests bias from misspecifying "
        "tail-heaviness, not optimizer convergence to the true nu (a "
        "different question from the nu\\_mode=\"learned\" single-point "
        "diagnostic in Table C1). alpha\\_ref/omega\\_ref: arch's own "
        "ARCH(1)-t MLE for that series (never re-estimated here). Verdict: "
        "PASS iff relative error on alpha (and beta, if the GARCH(1,1) "
        "extension) is below the 10% tolerance used throughout this project. "
        "QLIKE/LL\\_t evaluated OOS on the test split using the fixed nu of "
        "that row (not a learned nu\\_hat, which does not exist in "
        "nu\\_mode=\"fixed\")."
    )
    stem = tables_dir / "TableC2_arch_restricted_lambda_nu_sensitivity"
    _save_all(df, stem, "Table C2. ARCH(1)-Restricted LSTM -- ($\\lambda$, $\\nu$) Sensitivity Sweep",
              "tab:arch_restricted_lambda_nu_sensitivity", note)


# ──────────────────────────────────────────────────────────────────────────────
# Table C3 — ARCH(1) vs. ARCH-LSTM head-to-head (OOS metrics + VaR 99% backtest)
# ──────────────────────────────────────────────────────────────────────────────

def build_table_c3_arch1_vs_archlstm(all_results: dict, series_list: list[str], out_dir: Path) -> None:
    """
    Table C3: side-by-side comparison of ARCH(1) (traditional, arch's own
    MLE), GARCH(1,1) (traditional, the project's canonical benchmark),
    ARCH-LSTM (the ARCH(1)-restricted architecture's OOS predictions, mean
    over seeds), and -- once trained -- GARCH-LSTM (the GARCH(1,1)-
    restricted extension, same OOS convention) -- one row per series,
    each model's point-forecast metrics and its 99% VaR/ES backtest
    outcome. Pulls directly from raw_results.json (src.eval.run_all_
    metrics); requires the 3 REQUIRED models to be present for a given
    series (silently skips otherwise).
    GARCH-LSTM is OPTIONAL: its column block is only added if at least
    one series has it in raw_results.json (i.e. this table upgrades
    itself to 4 models automatically once that model's official run +
    eval pipeline have populated raw_results.json -- no code change
    needed here when that happens).
    """
    required_models = ["ARCH(1)", "GARCH(1,1)", "ARCH-LSTM"]
    optional_models = ["GARCH-LSTM"]
    models = required_models + [
        m for m in optional_models
        if any(m in all_results.get(s, {}) for s in series_list)
    ]
    metric_cols = ["MSE", "MAE", "R2", "QLIKE"]
    var_level = "0.99"
    col_tuples = [(m, c) for m in models for c in metric_cols] + \
                 [(m, c) for m in models for c in
                  ["VaR exc/T (99%)", "Kupiec p_uc", "KupiecPass", "Christ. p_cc", "ChristPass", "ES Z2 (99%)"]]
    cols = pd.MultiIndex.from_tuples(col_tuples)

    rows = []
    kept_series = []
    for series in series_list:
        s_res = all_results.get(series, {})
        if not all(m in s_res for m in required_models):
            missing = [m for m in required_models if m not in s_res]
            log.warning("[%s] missing %s in raw_results.json — "
                        "skipping Table C3 row.", series, ", ".join(missing))
            continue
        kept_series.append(series)
        row = []
        for m in models:
            metrics = s_res.get(m, {}).get("metrics", {})
            for c in metric_cols:
                v = metrics.get(c)
                row.append(_fmt(v, 4) if c != "R2" else _fmt(v, 4))
        for m in models:
            ve = s_res.get(m, {}).get("var_es", {}).get(var_level, {})
            st = ve.get("student_t", {})
            es = ve.get("es_backtest", {})
            n_exc, T = st.get("n_exc"), st.get("T")
            exc_str = f"{n_exc}/{T}" if n_exc is not None and T is not None else "—"
            row.extend([
                exc_str,
                _fmt(st.get("p_uc"), 4),
                _fmt_bool_pass(st.get("p_uc")),
                _fmt(st.get("p_cc"), 4),
                _fmt_bool_pass_cc(st.get("p_cc")),
                _fmt(es.get("Z2"), 4),
            ])
        rows.append(row)

    if not rows:
        log.warning("No series had all of %s — skipping Table C3.", ", ".join(required_models))
        return

    df = pd.DataFrame(rows, index=kept_series, columns=cols)
    garch11_note = (
        " GARCH-LSTM: the GARCH(1,1)-restricted extension "
        "(trainable forget gate, bounded persistence/mix reparametrization, "
        "same S=10-seed OOS convention) -- see Table C1 for whether it "
        "recovers GARCH(1,1)'s own (alpha, beta) parameters."
        if "GARCH-LSTM" in models else ""
    )
    note = (
        "ARCH(1) and GARCH(1,1): arch package's own Student-t MLE "
        "(traditional econometric estimators; GARCH(1,1) is this project's "
        "canonical benchmark). ARCH-LSTM: the ARCH(1)-restricted LSTM cell "
        "(src.models.arch\\_restricted), mean OOS forecast over S=10 seeds "
        "at (lambda=1.0, nu learned) -- see Table C1 for whether it "
        "actually recovers ARCH(1)'s own parameters (it does not, in any "
        "of the 6 series as of this run)." + garch11_note +
        " QLIKE = T\\textsuperscript{-1}"
        "Sigma[ln sigma\\textsuperscript{2}\\_t + eps\\textsuperscript{2}\\_t/"
        "sigma\\textsuperscript{2}\\_t]. VaR/ES: Student-t 99%% backtest "
        "(src.eval.var\\_es\\_backtest); KupiecPass/ChristPass = Yes when "
        "p >= 0.05. ES Z2: Acerbi-Szekely (2014) statistic (src.eval."
        "var\\_es\\_backtest); its own p-value is in the underlying JSON, "
        "not shown here for space -- see Table 11/12 for the full backtest."
    )
    stem = out_dir / "TableC3_arch1_vs_archlstm"
    _save_all(df, stem, "Table C3. ARCH(1) vs. GARCH(1,1) vs. ARCH-LSTM -- OOS Metrics and VaR(99%) Backtest",
              "tab:arch1_vs_archlstm", note)


# ──────────────────────────────────────────────────────────────────────────────
# Table C4 — ARCH-LSTM forecast metrics by split (train / validation / test)
# ──────────────────────────────────────────────────────────────────────────────

def build_table_c4_archlstm_by_split(tables_dir: Path) -> None:
    """
    Table C4: ARCH-LSTM's own MSE/MAE/R2/QLIKE, reported separately for
    the train, validation, and test/OOS windows -- unlike Table C1/C3,
    which only ever show the test split. Source: outputs/tables/
    archlstm_by_split_raw.json (src.eval.arch_restricted_recovery.
    compute_archlstm_by_split --by-split), which re-predicts sigma2 from
    each series' already-saved seed weights (no retraining).
    """
    src_path = tables_dir / "archlstm_by_split_raw.json"
    if not src_path.exists():
        log.warning("No %s found — skipping Table C4 "
                     "(run `python -m src.eval.arch_restricted_recovery --by-split` first).",
                     src_path)
        return
    by_series = json.loads(src_path.read_text())
    if not by_series:
        log.warning("Empty %s — skipping Table C4.", src_path)
        return

    splits = ["train", "validation", "test"]
    metric_cols = ["MSE", "MAE", "R2", "QLIKE"]
    col_tuples = [(s.capitalize(), c) for s in splits for c in metric_cols]
    cols = pd.MultiIndex.from_tuples(col_tuples)

    rows = []
    kept_series = []
    for series, split_data in by_series.items():
        kept_series.append(series)
        row = []
        for split in splits:
            m = split_data.get(split, {})
            for c in metric_cols:
                row.append(_fmt(m.get(c), 4))
        rows.append(row)

    df = pd.DataFrame(rows, index=kept_series, columns=cols)
    note = (
        "ARCH-LSTM (src.models.arch\\_restricted), mean OOS/in-sample "
        "forecast over S=10 seeds, re-predicted from each seed's saved "
        "weights (no retraining) on the train, validation, and test "
        "windows separately -- Table C1/C3 report the test split only. "
        "Train/validation MSE/MAE are in-sample-adjacent (the model saw "
        "these observations during training via the val\\_loss early-"
        "stopping criterion for validation, and directly for train), so "
        "they are NOT forecast-accuracy claims in the same sense as the "
        "test column -- included here purely to show the split-to-split "
        "gap this restricted architecture exhibits. "
        "QLIKE = T\\textsuperscript{-1}Sigma[ln sigma\\textsuperscript{2}\\_t "
        "+ eps\\textsuperscript{2}\\_t/sigma\\textsuperscript{2}\\_t]."
    )
    stem = tables_dir / "TableC4_archlstm_by_split"
    _save_all(df, stem, "Table C4. ARCH-LSTM -- Forecast Metrics by Split (Train / Validation / Test)",
              "tab:archlstm_by_split", note)


# ──────────────────────────────────────────────────────────────────────────────
# Main
# ──────────────────────────────────────────────────────────────────────────────

def run(config_path: str) -> None:
    cfg           = load_config(config_path)
    tables_dir    = Path(cfg["paths"]["tables"])
    models_dir    = Path(cfg["paths"]["models"])
    processed_dir = Path(cfg["paths"]["processed_data"])
    tables_dir.mkdir(parents=True, exist_ok=True)

    raw_path = tables_dir / "raw_results.json"
    if not raw_path.exists():
        log.error("raw_results.json not found — run eval step first.")
        return
    all_results = json.loads(raw_path.read_text())

    enc_path = tables_dir / "encompassing_results.json"
    encompassing_all = json.loads(enc_path.read_text()) if enc_path.exists() else {}

    deg_path = tables_dir / "degeneracy_flags.json"
    degeneracy_all = json.loads(deg_path.read_text()) if deg_path.exists() else {}
    if not deg_path.exists():
        log.warning("degeneracy_flags.json not found — Tables 4-7 DEGENERATE column "
                     "will show '—' (run `python -m src.eval.degeneracy` first).")

    series_list = [s["name"] for s in cfg["series"]]

    # ── Table 3: roster ───────────────────────────────────────────────────────
    log.info("Building Table 3 (roster) …")
    build_table3(tables_dir)

    # ── Tables 4–7: OOS performance ───────────────────────────────────────────
    for i, series in enumerate(series_list, start=4):
        log.info("Building Table %d (OOS %s) …", i, series)
        build_table_oos(series, all_results.get(series, {}), tnum=i, out_dir=tables_dir,
                         degeneracy=degeneracy_all.get(series, {}))

    # ── Table 8: cross-market summary ─────────────────────────────────────────
    log.info("Building Table 8 (cross-market) …")
    build_table8(all_results, series_list, tables_dir)

    # ── Table 9: Diebold–Mariano ──────────────────────────────────────────────
    log.info("Building Table 9 (DM) …")
    build_table9(all_results, series_list, tables_dir)

    # ── Table 10: Forecast-Encompassing ───────────────────────────────────────
    if encompassing_all:
        log.info("Building Table 10 (Encompassing) …")
        build_table10(encompassing_all, series_list, tables_dir)
    else:
        log.warning("encompassing_results.json not found — skipping Table 10.")

    # ── Table 11a–11d: Risk backtests ───────────────────────────────────────
    for tlabel, series in zip(["a", "b", "c", "d"], series_list):
        log.info("Building Table 11%s (Risk %s) …", tlabel, series)
        build_table11_risk(series, all_results.get(series, {}), tlabel=tlabel, out_dir=tables_dir)

    # ── Table 12: risk backtests by market (Section 9.5) ──────────────────────
    log.info("Building Table 12 (risk backtests by market) …")
    build_table12_risk_summary(all_results, series_list, tables_dir)

    # ── Table 4e: GARCH estimation ────────────────────────────────────────────
    log.info("Building Table 4e (GARCH estimation) …")
    build_table4_estimation(cfg, models_dir, processed_dir, tables_dir)

    # ── Tables A1–A4: full estimation ─────────────────────────────────────────
    for idx, series in enumerate(series_list, start=1):
        log.info("Building Table A%d (estimation %s) …", idx, series)
        build_table_Ax(series, idx, models_dir, tables_dir)

    # ── Table 13: gate <-> GARCH parameter correspondence (Section 9.3) ──────
    log.info("Building Table 13 (gate correspondence) …")
    build_table13_gate_correspondence(tables_dir)

    # ── Table B1: ablation ladder (Section 7 / Proposition 2) ────────────────
    log.info("Building Table B1 (ablation ladder) …")
    build_table_b1(tables_dir)

    # ── Table C2: ARCH(1)-restricted LSTM (lambda, nu) sensitivity sweep ─────
    log.info("Building Table C2 (ARCH-LSTM lambda/nu sensitivity) …")
    build_table_c2_lambda_nu_sensitivity(tables_dir, models_dir)

    # ── Table C3: ARCH(1) vs. ARCH-LSTM head-to-head (metrics + VaR) ─────────
    log.info("Building Table C3 (ARCH(1) vs ARCH-LSTM) …")
    build_table_c3_arch1_vs_archlstm(all_results, series_list, tables_dir)

    # ── Table C4: ARCH-LSTM forecast metrics by split ─────────────────────────
    log.info("Building Table C4 (ARCH-LSTM by split) …")
    build_table_c4_archlstm_by_split(tables_dir)

    log.info("══ build_tables complete — outputs in %s ══", tables_dir)


def main():
    parser = argparse.ArgumentParser(description="Build all paper tables")
    parser.add_argument("--config", default="config/config.yaml")
    args = parser.parse_args()
    run(args.config)


if __name__ == "__main__":
    import argparse
    main()
